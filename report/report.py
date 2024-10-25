from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# Create your views here.

def doc_write(request):
    # 워드 문서 생성
    doc = Document()

    # 제목 추가
    title = doc.add_heading('제목을 이곳에 작성합니다', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 단락 추가
    p = doc.add_paragraph('첫번째 단락입니다')
    # 두 번째, 세 번째 단락을 추가
    p2 = doc.add_paragraph('이것은 두 번째 단락입니다.')
    p3 = doc.add_paragraph('이것은 세 번째 단락입니다.')

    # 첫 번째 단락 앞에 새 단락 끼워넣기
    p.insert_paragraph_before('이것은 첫 번째 단락 앞에 새로 끼워넣은 단락입니다.')

    # 두 번째 단락 텍스트 변경
    p2.clear()
    p2.add_run('두 번째 단락의 텍스트가 변경되었습니다.')

    # 세 번째 단락 텍스트 변경
    p3.clear()
    p3.add_run('세 번째 단락의 텍스트가 변경되었습니다.')

    # 3x3 크기의 표 추가
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid' # 표의 각 셀에 격자선 추가

    # 첫 번째 행의 셀들을 hdr_cells 변수에 할당하고 각 셀에 헤더 데이터 추가
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '헤더 1'
    hdr_cells[1].text = '헤더 2'
    hdr_cells[2].text = '헤더 3'

    # for문으로 각 행의 셀에 접근하여 두 번째와 세 번째 행에 데이터 추가
    for i in range(1, 3):
        row_cells = table.rows[i].cells
        row_cells[0].text = f'행 {i}, 열 1'
        row_cells[1].text = f'행 {i}, 열 2'
        row_cells[2].text = f'행 {i}, 열 3'

    # 리스트 추가
    doc.add_paragraph(
        '첫 번째 항목', style='List Number'
    )
    doc.add_paragraph(
        '두 번째 항목', style='List Number'
    )

    doc.add_picture('testplot.png', width=Inches(1.5))
    doc.save('example_image_001.docx')
    doc = aw.Document('example_image_001.docx')
    doc.save('./static/images/exampleimage003.png')

    return render(request, 'report/report_list.html', context)


def report_write(request):
    # csv파일을 쓰기모드 'w'로 열기(인코딩 'cp949', 라인 종료 문자 설정 'newline = '')
    with open('example.csv', 'w', encoding='cp949', newline='') as file:
        csv_writer = csv.writer(file)  # 파일 객체를 csv.writer의 인자로 전달해 새로운 writer 객체를 생성
        csv_writer.writerow(['이름', '나이', '직업'])  # 헤더 작성
        csv_writer.writerow(['박은영', 30, '엔지니어'])  # 데이터 행 추가
        csv_writer.writerow(['김세빛', 25, '디자이너'])
        csv_writer.writerow(['안희수', 35, '의사'])
        csv_writer.writerow(['정현욱', 40, '선생님'])
        csv_writer.writerow(['강찬영', 22, '학생'])

        context ={
    }

    return render(request, 'report/report_list.html', context)